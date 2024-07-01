from dataclasses import dataclass
from docx import Document

from ._base import RawDoc


@dataclass
class DocxLoader(RawDoc):
    def extract_text(self):
        doc = Document(self.file_path)
        for paragraph in doc.paragraphs:
            if paragraph.text:
                yield paragraph.text
            else:
                continue

    def extract_image(self):
        doc = Document(self.file_path)
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if run.text:
                    continue
                else:
                    for inline in run.element.iter():  # type: ignore
                        if inline.tag.endswith("inline"):  # type: ignore
                            for pic in inline.iter():  # type: ignore
                                if pic.tag.endswith("blip"):  # type: ignore
                                    image = pic.embed  # type: ignore
                                    image_part = run.part.related_parts[image]
                                    yield image_part.blob
                                else:
                                    continue
                        else:
                            continue
