import base64
import io
import tempfile
from dataclasses import dataclass
from pathlib import Path
from typing import Callable, Generator, Literal, Union, NamedTuple

from fastapi import FastAPI, UploadFile
from fastapi.responses import StreamingResponse
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from requests import get

class RawDoc(NamedTuple):
	type: Literal["file", "url", "upload", "io", "bytes"]
	raw_doc: Union[str, Path, UploadFile, io.BytesIO, bytes]


def pt_to_px(pt: float) -> float:
	return pt * 1.33333


def get_font_color(run: object) -> str:
	try:
		if run.font.color.rgb:
			r, g, b = run.font.color.rgb
			return f"#{r:02x}{g:02x}{b:02x}"
	except AttributeError:
		pass
	return "#000000"  # Default color black if no color is set


def get_paragraph_alignment(paragraph) -> str:
	if paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER:
		return "center"
	elif paragraph.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
		return "right"
	return "left"  # Default to left alignment

def load_images(run):
	for inline in run.element.iter():
		if inline.tag.endswith("inline"):
			for pic in inline.iter():
				if pic.tag.endswith("blip"):
					image = pic.embed
					image_part = run.part.related_parts[image]
					image_stream = image_part.blob
					image_base64 = base64.b64encode(image_stream).decode("utf-8")
					image_type = image_part.content_type
					yield f"<img src='data:{image_type};base64,{image_base64}' class='mx-auto' />"


def use_tempfile(*,raw_data:RawDoc) -> str:
	with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as f:
		if raw_data.type == "file" and isinstance(raw_data.raw_doc, Path):
			f.write(raw_data.raw_doc.read_bytes())
		elif raw_data.type == "file" and isinstance(raw_data.raw_doc, str):
			f.write(Path(raw_data.raw_doc).read_bytes())
		elif raw_data.type == "upload":
			assert isinstance(raw_data.raw_doc, UploadFile)
			f.write(raw_data.raw_doc.file.read())
		elif raw_data.type == "io":
			assert isinstance(raw_data.raw_doc, io.BytesIO)
			f.write(raw_data.raw_doc.getvalue())
		elif raw_data.type == "bytes":
			assert isinstance(raw_data.raw_doc, bytes)
			f.write(raw_data.raw_doc)
		elif raw_data.type == "url":
			assert isinstance(raw_data.raw_doc, str)
			response = get(raw_data.raw_doc)
			f.write(response.content)
		else:
			raise ValueError(f"Invalid source type: {raw_data.type}")
		return f.name
		

def stream_docx(path: str) -> Generator[str, None, None]:
	doc = Document(path)
	yield "<!DOCTYPE html>"
	yield "<html>"
	yield """<head>
	<script src='https://cdn.tailwindcss.com'></script>
	<style>
	body { font-family: Arial, sans-serif; line-height: 1.6; padding: 20px; }
	img { max-width: 100%; height: auto; }
	</style>
	</head>"""
	yield "<body>"

	for element in doc.element.body:
		if element.tag.endswith("p"):
			paragraph = doc.paragraphs[doc.element.body.index(element)]
			alignment = get_paragraph_alignment(paragraph)
			yield f"<p style='text-align: {alignment};'>"
			for run in paragraph.runs:
				if run.text:
					# This is a text run
					text = run.text
					font_size = pt_to_px(run.font.size.pt) if run.font.size else 16
					font_color = get_font_color(run)
					bold = "font-bold" if run.bold else ""
					italic = "italic" if run.italic else ""
					underline = "underline" if run.underline else ""
					yield f"<span style='font-size:{font_size}px; color:{font_color};' class='{bold} {italic} {underline}'>{text}</span>"
				else:
					# This is an image run
					for image in load_images(run):
						yield image
			yield "</p>"
		elif element.tag.endswith("tbl"):
			table = doc.tables[doc.element.body.index(element)]
			yield "<table>"
			for row in table.rows:
				yield "<tr>"
				for cell in row.cells:
					yield "<td>"
					for paragraph in cell.paragraphs:
						for run in paragraph.runs:
							if run.text:
								# This is a text run
								text = run.text
								font_size = pt_to_px(run.font.size.pt) if run.font.size else 16
								font_color = get_font_color(run)
								bold = "font-bold" if run.bold else ""
								italic = "italic" if run.italic else ""
								underline = "underline" if run.underline else ""
								yield f"<span style='font-size:{font_size}px; color:{font_color};' class='{bold} {italic} {underline}'>{text}</span>"
							else:
								# This is an image run
								for image in load_images(run):
									yield image
					yield "</td>"
				yield "</tr>"
			yield "</table>"
	yield "</body>"
	yield "</html>"


@dataclass
class StreamDOCX:
	path: Union[str, Path, UploadFile, io.BytesIO, bytes]
	source_type: Literal["file", "url", "upload", "io", "bytes"]

	def handler(
		self,
		source: Union[str, Path, UploadFile, io.BytesIO, bytes],
		callback: Callable[[str], Generator[str, None, None]],
	) -> StreamingResponse:
		
		return StreamingResponse(callback(self.path), media_type="text/html")

	def run(self) -> StreamingResponse:
		return self.handler(self.path, stream_docx)
